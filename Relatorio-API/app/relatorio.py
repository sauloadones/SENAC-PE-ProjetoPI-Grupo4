import sqlite3
import pandas as pd
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

def gerar_pdf(db_file):
    db_file.save('temp.db')
    conn = sqlite3.connect('temp.db')
    df = pd.read_sql_query("SELECT * FROM sua_tabela", conn)

    pdf_file = 'relatorio.pdf'
    c = canvas.Canvas(pdf_file, pagesize=letter)
    width, height = letter
    textobject = c.beginText(40, height - 50)
    textobject.setFont("Helvetica", 10)

    textobject.textLine(" | ".join(df.columns))
    textobject.textLine("-" * 100)

    for _, row in df.iterrows():
        line = " | ".join(str(val) for val in row)
        textobject.textLine(line)
        if textobject.getY() < 50:
            c.drawText(textobject)
            c.showPage()
            textobject = c.beginText(40, height - 50)
            textobject.setFont("Helvetica", 10)

    c.drawText(textobject)
    c.showPage()
    c.save()
    return pdf_file

def gerar_excel(db_file):
    db_file.save('temp.db')
    conn = sqlite3.connect('temp.db')
    df = pd.read_sql_query("SELECT * FROM sua_tabela", conn)

    output_excel = 'relatorio.xlsx'
    df.to_excel(output_excel, index=False)
    return output_excel

def gerar_word(db_file):
    db_file.save('temp.db')
    conn = sqlite3.connect('temp.db')
    df = pd.read_sql_query("SELECT * FROM sua_tabela", conn)

    doc = Document()
    doc.add_heading('Relatório', 0)

    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col

    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val)

    output_docx = 'relatorio.docx'
    doc.save(output_docx)
    return output_docx
